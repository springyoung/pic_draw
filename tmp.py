#coding:utf-8
import dateToPic as dt
#import pywt
import numpy as np
import matplotlib.pyplot as plt
import sys
from docx import Document
from docx.shared import Inches
#reload(sys)
#sys.setdefaultencoding('utf8')
count = 0
document = Document()
p = document.add_paragraph('附录')
pattern = {'speed':['1节','2节','3节','4节','5节','6节'],'point':['测点1','测点2','测点3', \
                 '测点4','测点5','测点6','测点7','测点8','测点9','测点10','测点11', \
                 '测点12','测点13','测点14','测点15','测点16'], \
                 'geshan':['优化','实船','无'],'famen':['全开','半开','全关']}
cedian_act = ['加速度测点1','加速度测点2','加速度测点3','加速度测点4','加速度测点5','加速度测点6','加速度测点7','加速度测点8', \
                '压力测点1','压力测点2','压力测点3','压力测点4','压力测点5','压力测点6','压力测点7','压力测点8']
for cedian_index,point in enumerate(pattern['point']):
    for geshan in pattern['geshan']:
        for speed in pattern['speed']:
            for famen in pattern['famen']:
                dt.plot_by_re([point,speed,famen,geshan],log=False,lvbo = 0,file_path='./' + geshan + '格栅' + speed + '阀门' + famen + cedian_act[cedian_index] +'.png',pic_title=geshan + '格栅' + speed + '阀门' + famen + cedian_act[cedian_index])
                count += 1
                document.add_picture('./' + geshan + '格栅' + speed + '阀门' + famen + cedian_act[cedian_index] +'.png',width = Inches(6))
                document.add_paragraph('附图' + str(count) + ': ' + geshan + '格栅' + speed + '阀门' + famen + cedian_act[cedian_index])
document.save('./futu_new.docx')

#count = 0
#document = Document()
# p = document.add_paragraph('阀门开启状态对低频线谱的影响')
# p = document.add_paragraph('实船格栅为例')
# for speed in pattern['speed']:
#     for cedian in ['测点9','测点10','测点15','测点16']:
#         dt.plot_by_re([speed,'实船','全开|半开|全关',cedian],log=False,lvbo = 0,file_path='./' + str(count) + 'udf.png',pic_title='')
#         document.add_picture('./' + str(count) + 'udf.png',width = Inches(6))
#         count += 1
# p = document.add_paragraph('不同航速对比')
# p = document.add_paragraph('阀门状态全开，对比航速3节，6节的频谱')
# p = document.add_paragraph('压力测点')
# for geshan in pattern['geshan']:
#     p = document.add_paragraph('格栅状态' +geshan)
    
#     for cedian in ['测点9','测点10','测点15','测点16']:
#         dt.plot_by_re([geshan,'全开','3节|6节',cedian],log=False,lvbo = 0,file_path='./' + str(count) + 'udf.png',pic_title='')
#         document.add_picture('./' + str(count) + 'udf.png',width = Inches(6))
#         count += 1

# #  dt.plot_by_re([geshan,'全开','3节|6节',cedian],log=False,lvbo = 0,file_path='./udf.png',pic_title='')
# p = document.add_paragraph('加速度测点')
# for geshan in pattern['geshan']:
#     p = document.add_paragraph('格栅状态' +geshan)
    
#     for cedian in ['测点1','测点2','测点6','测点7']:
#         dt.plot_by_re([geshan,'全开','3节|6节',cedian],log=False,lvbo = 0,file_path='./' + str(count) + 'udf.png',pic_title='')
#         document.add_picture('./' + str(count) + 'udf.png',width = Inches(6))
#         count += 1

# p = document.add_paragraph('实船格栅与优化格栅对比')
# p = document.add_paragraph('阀门状态全开')
# for speed in ['6节']:
#     p = document.add_paragraph('航速为' + speed +'时')
#     p = document.add_paragraph('压力测点')
#     for cedian in ['测点9','测点10','测点15','测点16']:
#         dt.plot_by_re(['实|优化','全开',speed,cedian],log=False,lvbo = 0,file_path='./' + str(count) + 'udf.png',pic_title='')
#         document.add_picture('./' + str(count) + 'udf.png',width = Inches(6))
#         count += 1
#     p = document.add_paragraph('加速度测点')
#     for cedian in ['测点1','测点2','测点6','测点7']:
#         dt.plot_by_re(['实|无','全开',speed,cedian],log=False,lvbo = 0,file_path='./' + str(count) + 'udf.png',pic_title='')
#         document.add_picture('./' + str(count) + 'udf.png',width = Inches(6))
#         count += 1

# p = document.add_paragraph('无格栅与优化格栅对比')
# p = document.add_paragraph('阀门状态全开')
# for speed in ['3节','6节']:
#     p = document.add_paragraph('航速为' + speed +'时')
#     p = document.add_paragraph('压力测点')
#     for cedian in ['测点9','测点10','测点15','测点16']:
#         dt.plot_by_re(['优化|无','全开',speed,cedian],log=False,lvbo = 0,file_path='./' + str(count) + 'udf.png',pic_title='')
#         document.add_picture('./' + str(count) + 'udf.png',width = Inches(6))
#         count += 1
#     p = document.add_paragraph('加速度测点')
#     for cedian in ['测点1','测点2','测点6','测点7']:
#         dt.plot_by_re(['优化|无','全开',speed,cedian],log=False,lvbo = 0,file_path='./' + str(count) + 'udf.png',pic_title='')
#         document.add_picture('./' + str(count) + 'udf.png',width = Inches(6))
#         count += 1

# p = document.add_paragraph('测点对应的加速度和压力脉动数据对比')
# p = document.add_paragraph('航速为3节，阀门状态为全开')
# cedian_num1 = [1,2,6,7]
# cedian_num2 = [9,10,15,16]
# for geshan in pattern['geshan']:
#     for index,i in enumerate(cedian_num1):
#         dt.plot_by_re([geshan,'全开','3节','测点' + str(i)],log=False,lvbo = 0,file_path='./' + str(count) + 'udf.png',pic_title='')
#         document.add_picture('./' + str(count) + 'udf.png',width = Inches(6))
#         count += 1
#         dt.plot_by_re([geshan,'全开','3节','测点' + str(cedian_num2[index])],log=False,lvbo = 0,file_path='./' + str(count) + 'udf.png',pic_title='')
#         document.add_picture('./' + str(count) + 'udf.png',width = Inches(6))
#         count += 1


#document.save('./2.docx')

